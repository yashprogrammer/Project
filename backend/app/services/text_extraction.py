"""
Text extraction service for various document formats.
Supports: .txt, .md, .pdf, .docx
"""
import os
from pypdf import PdfReader
from docx import Document as DocxDocument


class TextExtractionService:
    """Service to extract text from various document formats"""
    
    SUPPORTED_FORMATS = {
        'text/plain': ['txt', 'md'],
        'application/pdf': ['pdf'],
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['docx'],
    }
    
    def is_supported(self, content_type: str, file_path: str) -> bool:
        """Check if the file format is supported"""
        extension = self._get_extension(file_path)
        
        # Check by content type
        if content_type in self.SUPPORTED_FORMATS:
            return True
        
        # Check by extension
        for mime_type, extensions in self.SUPPORTED_FORMATS.items():
            if extension in extensions:
                return True
        
        return False
    
    def extract_text(self, file_path: str, content_type: str) -> str:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the file
            content_type: MIME type of the file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
            Exception: For other extraction errors
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = self._get_extension(file_path)
        
        # Plain text files (.txt, .md)
        if content_type == 'text/plain' or extension in ['txt', 'md']:
            return self._extract_text_file(file_path)
        
        # PDF files
        elif content_type == 'application/pdf' or extension == 'pdf':
            return self._extract_pdf(file_path)
        
        # Word documents (.docx)
        elif 'wordprocessingml' in content_type or extension == 'docx':
            return self._extract_docx(file_path)
        
        else:
            raise ValueError(
                f"Unsupported file format: {content_type} (extension: {extension}). "
                f"Supported formats: .txt, .md, .pdf, .docx"
            )
    
    def _get_extension(self, file_path: str) -> str:
        """Get file extension without the dot"""
        return os.path.splitext(file_path)[1].lstrip('.').lower()
    
    def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text.strip()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()
            return text.strip()
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF files using pypdf"""
        try:
            reader = PdfReader(file_path)
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            full_text = '\n\n'.join(text_parts)
            return full_text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from Word documents (.docx)"""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = '\n\n'.join(text_parts)
            return full_text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

